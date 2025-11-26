# 3_rewriter_template.py (Enhanced)
import os
from docx import Document
from sentence_transformers import SentenceTransformer, util
import torch
import ast


def extract_section(file_path, section_name):
    """Extract a list from parsed NER output file."""
    with open(file_path, "r", encoding="utf-8") as f:
        for line in f:
            if line.startswith(f"{section_name}:"):
                try:
                    return ast.literal_eval(line.split(":", 1)[1].strip())
                except Exception:
                    return []
    return []


def generate_ats_resume(cv_data_path, jd_data_path, output_dir):
    # --- Extract NER data ---
    cv_skills = extract_section(cv_data_path, "SKILLS")
    jd_skills = extract_section(jd_data_path, "SKILLS")
    cv_exp = extract_section(cv_data_path, "EXPERIENCE")
    cv_projects = extract_section(cv_data_path, "PROJECTS")
    cv_education = extract_section(cv_data_path, "EDUCATION")
    cv_name = extract_section(cv_data_path, "NAME")

    # --- Initialize model ---
    model = SentenceTransformer("all-MiniLM-L6-v2")

    # --- Compute skill matches ---
    if cv_skills and jd_skills:
        cv_emb = model.encode(cv_skills, convert_to_tensor=True)
        jd_emb = model.encode(jd_skills, convert_to_tensor=True)
        cosine_scores = util.cos_sim(cv_emb, jd_emb)

        # Keep matched skills above similarity threshold
        matched_skills = [
            cv_skills[i]
            for i in range(len(cv_skills))
            if torch.max(cosine_scores[i]) > 0.55
        ]
    else:
        cosine_scores = torch.zeros(1)
        matched_skills = cv_skills

    missing_skills = [s for s in jd_skills if s not in matched_skills]

    # --- Compute experience relevance ---
    exp_relevance = []
    if cv_exp and jd_skills:
        exp_emb = model.encode(cv_exp, convert_to_tensor=True)
        jd_emb = model.encode(jd_skills, convert_to_tensor=True)
        exp_sim = util.cos_sim(exp_emb, jd_emb)

        # Relevance = average cosine similarity with JD skills
        for i in range(len(cv_exp)):
            score = torch.mean(exp_sim[i]).item()
            exp_relevance.append((cv_exp[i], score))

        # Sort experiences by relevance (descending)
        exp_relevance.sort(key=lambda x: x[1], reverse=True)
    else:
        exp_relevance = [(e, 0.0) for e in cv_exp]

    # --- Rewrite experience sentences emphasizing matched skills ---
    rewritten_exp = []
    for exp, score in exp_relevance:
        relevant_skills = [s for s in matched_skills if s.lower() in exp.lower()]
        new_sentence = exp
        if relevant_skills:
            new_sentence += f" (Key Skills: {', '.join(relevant_skills)})"
        elif score > 0.4:
            new_sentence += " (Aligned with Job Role Requirements)"
        rewritten_exp.append(new_sentence)

    # --- Generate DOCX resume ---
    doc = Document()

    # --- Header ---
    if cv_name:
        doc.add_heading(cv_name[0], level=0)
    else:
        doc.add_heading("Candidate Name", level=0)
    doc.add_paragraph("Email: [your email] | Phone: [your phone] | LinkedIn: [your linkedin]")

    # --- Professional Summary ---
    doc.add_heading("Professional Summary", level=1)
    summary_text = (
        f"Results-driven professional with expertise in {', '.join(matched_skills[:5])}. "
        f"Demonstrated success in roles related to {', '.join(jd_skills[:3])}. "
        f"Actively enhancing proficiency in {', '.join(missing_skills[:3])} to meet evolving industry demands."
    )
    doc.add_paragraph(summary_text)

    # --- Skills Section ---
    doc.add_heading("Key Skills", level=1)
    doc.add_paragraph(", ".join(sorted(set(matched_skills))))

    # --- Recommended Skills ---
    if missing_skills:
        doc.add_heading("Recommended Skills to Learn", level=1)
        doc.add_paragraph(", ".join(missing_skills))

    # --- Experience ---
    doc.add_heading("Experience", level=1)
    for exp in rewritten_exp:
        doc.add_paragraph(exp, style="List Bullet")

    # --- Projects ---
    if cv_projects:
        doc.add_heading("Projects", level=1)
        for proj in cv_projects:
            doc.add_paragraph(proj, style="List Bullet")

    # --- Education ---
    doc.add_heading("Education", level=1)
    for edu in cv_education:
        doc.add_paragraph(edu, style="List Bullet")

    # --- Additional Info ---
    doc.add_heading("Additional Information", level=1)
    doc.add_paragraph("Certifications, Achievements, or Interests can be added here.")

    # --- Save the file in the specified directory ---
    os.makedirs(output_dir, exist_ok=True)
    save_path = os.path.join(output_dir, "ATS_Optimized_Resume.docx")
    doc.save(save_path)
    print(f"✅ ATS Resume saved successfully at:\n{save_path}")


# --- Run ---
if __name__ == "__main__":
    generate_ats_resume(
        "ML/CV/CV_NER_output.txt",
        "ML/Job_Description/JD_NER_output.txt",
        r"C:\Users\HP\OneDrive\Desktop\Major"
    )
